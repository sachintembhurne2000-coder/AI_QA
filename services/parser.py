"""
Parser Service — unified document ingestion for PDF, DOCX, TXT, YAML/JSON
"""
import io
import re
from pathlib import Path
from typing import Union


def parse_document(file_obj, filename: str) -> str:
    """
    Accept a file-like object or bytes and return plain text.
    Supports: .pdf, .docx, .txt, .md, .yaml, .yml, .json, .log
    """
    name = filename.lower()
    content = file_obj.read() if hasattr(file_obj, "read") else file_obj

    if name.endswith(".pdf"):
        return _parse_pdf(content)
    elif name.endswith(".docx"):
        return _parse_docx(content)
    elif name.endswith((".txt", ".md", ".log", ".yaml", ".yml", ".json")):
        return _decode(content)
    else:
        # Attempt decode anyway
        return _decode(content)


def _decode(content: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(enc)
        except (UnicodeDecodeError, AttributeError):
            pass
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except ImportError:
        pass

    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts)
    except ImportError:
        pass

    return "[PDF parsing unavailable — install pdfplumber: pip install pdfplumber]"


def _parse_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        return "[DOCX parsing unavailable — install python-docx: pip install python-docx]"


def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for RAG ingestion."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
        i += chunk_size - overlap
    return chunks


def truncate_for_llm(text: str, max_chars: int = 60_000) -> str:
    """Truncate large texts, keeping head and tail for log analysis."""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + f"\n\n... [TRUNCATED — {len(text) - max_chars} chars omitted] ...\n\n"
        + text[-half:]
    )
