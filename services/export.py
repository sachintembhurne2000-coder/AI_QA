"""
Export Service — generate Excel, Word, JSON, and ZIP downloads
"""
import json
import io
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any


# ── Timestamp helper ────────────────────────────────────────────────────────
def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


# ── JSON export ─────────────────────────────────────────────────────────────
def to_json_bytes(data: Any) -> bytes:
    return json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")


# ── Excel export ─────────────────────────────────────────────────────────────
def test_cases_to_excel(test_cases: list[dict]) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    headers = ["ID", "Title", "Type", "Priority", "Preconditions",
               "Steps", "Expected Result", "Actual Result", "Status"]

    # Header styling
    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    ws.row_dimensions[1].height = 22

    # Data rows
    for row_idx, tc in enumerate(test_cases, 2):
        steps = tc.get("steps", [])
        steps_str = "\n".join(f"{i+1}. {s}" for i, s in enumerate(steps)) if isinstance(steps, list) else str(steps)
        values = [
            tc.get("id", f"TC-{row_idx-1:03d}"),
            tc.get("title", ""),
            tc.get("type", "Functional"),
            tc.get("priority", "Medium"),
            tc.get("preconditions", ""),
            steps_str,
            tc.get("expected_result", ""),
            "",  # Actual Result — blank for tester
            "Not Run",
        ]
        alt_fill = PatternFill("solid", fgColor="F7FAFC" if row_idx % 2 == 0 else "FFFFFF")
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col, value=val)
            cell.fill = alt_fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = border

    # Column widths
    widths = [10, 35, 14, 10, 25, 45, 35, 20, 10]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = w

    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Bug Report Word export ────────────────────────────────────────────────────
def bug_report_to_word(report: dict) -> bytes:
    """
    Generate a .docx bug report from a dict with keys:
    title, summary, severity, priority, environment,
    steps_to_reproduce, actual_result, expected_result,
    root_cause, logs_excerpt, recommendations
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Title
    title_para = doc.add_heading(report.get("title", "Bug Report"), level=1)
    title_para.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    def add_field(label: str, value: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(str(value))

    add_field("Severity", report.get("severity", "Medium"))
    add_field("Priority", report.get("priority", "P2"))
    add_field("Environment", report.get("environment", ""))
    doc.add_paragraph("")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(report.get("summary", ""))

    doc.add_heading("Steps to Reproduce", level=2)
    steps = report.get("steps_to_reproduce", [])
    if isinstance(steps, list):
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f"{i}. {step}", style="List Number")
    else:
        doc.add_paragraph(str(steps))

    doc.add_heading("Actual Result", level=2)
    doc.add_paragraph(report.get("actual_result", ""))

    doc.add_heading("Expected Result", level=2)
    doc.add_paragraph(report.get("expected_result", ""))

    doc.add_heading("Root Cause Analysis", level=2)
    doc.add_paragraph(report.get("root_cause", ""))

    if report.get("logs_excerpt"):
        doc.add_heading("Log Evidence", level=2)
        p = doc.add_paragraph(report["logs_excerpt"])
        p.runs[0].font.name = "Courier New"
        p.runs[0].font.size = Pt(9)

    doc.add_heading("Recommendations", level=2)
    recs = report.get("recommendations", [])
    if isinstance(recs, list):
        for r in recs:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph(str(recs))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── ZIP bundle ───────────────────────────────────────────────────────────────
def create_zip(files: dict[str, bytes]) -> bytes:
    """
    files: {filename: bytes}
    Returns ZIP bytes.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)
    return buf.getvalue()
